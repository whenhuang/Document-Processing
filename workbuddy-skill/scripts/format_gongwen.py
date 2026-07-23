#!/usr/bin/env python3
"""
公文格式排版脚本 - 按 GB/T 9704 党政机关公文格式标准排版 .docx 文档。

用法:
    python format_gongwen.py <输入文件.docx> [输出文件.docx]

若未指定输出文件，默认在输入文件名后加 "_已排版" 后缀。
"""

import sys
import re
import os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 格式常量 ──────────────────────────────────────────────

FONT_TITLE = "方正小标宋简体"       # 标题字体
SIZE_TITLE = Pt(22)                  # 二号

FONT_BODY = "仿宋_GB2312"            # 正文字体
SIZE_BODY = Pt(16)                   # 三号

FONT_HEITI = "黑体"                  # 一级标题字体
FONT_KAITI = "楷体_GB2312"           # 二级标题字体

# 表格（参照深圳港集团实际公文：仿宋小四、表头加粗、细单线边框）
FONT_TABLE = "仿宋_GB2312"           # 表格字体
SIZE_TABLE = Pt(14)                  # 小四（区别于正文三号）
TABLE_HEADER_BOLD = True             # 表头行加粗
TABLE_BORDER_SZ = 4                  # 边框 0.5pt（sz=4 即 1/8pt ×4）
TABLE_BORDER_COLOR = "000000"        # 边框黑色

# 发文字号 / 签发人 / 版记（上行文与版记，国标 GB/T 9704）
FONT_DOCNUM = "仿宋_GB2312"          # 发文字号
SIZE_DOCNUM = Pt(16)                 # 三号
FONT_SIGNER = "仿宋_GB2312"          # 签发人
SIZE_SIGNER = Pt(16)                 # 三号
FONT_BANJI = "仿宋_GB2312"           # 版记（抄送/印发）
SIZE_BANJI = Pt(12)                  # 四号

LINE_SPACING = Pt(28)                # 全文行距 28 磅（固定值）

# 标准页边距 (GB/T 9704，参照深圳港集团实际公文)
PAGE_TOP = Cm(3.7)
PAGE_BOTTOM = Cm(3.5)
PAGE_LEFT = Cm(2.8)
PAGE_RIGHT = Cm(2.6)

# 页脚页码边距
FOOTER_DISTANCE = Cm(2.5)

# 首行缩进（两个汉字 ≈ 32pt）
FIRST_INDENT = Pt(32)

# 页码字体（宋体四号 + 短横中点格式："— · 1 · —"）
FONT_PAGE_NUM = "宋体"
SIZE_PAGE_NUM = Pt(14)               # 四号

# ── 辅助函数 ──────────────────────────────────────────────


def set_font(run, name, size, bold=False, color=None):
    """设置 run 的字体属性，同时设置中文字体（eastAsia）"""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def set_line_spacing(paragraph, spacing_pt):
    """设置段落行距为固定值，段前段后归零"""
    pf = paragraph.paragraph_format
    pf.line_spacing = spacing_pt
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # 确保是精确行距（exactly），而非倍数
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(spacing_pt.pt * 20)))  # 转换为 twips
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')


def set_first_line_indent_chars(paragraph, chars):
    """设置首行缩进为指定字符数（同时写 firstLineChars + firstLine 双单位）

    chars: 字符数（如 2 表示 2 字符）
    同时显式写 leftChars=0/rightChars=0，避免 WPS 在"文本之前/之后"误显示。
    """
    pf = paragraph.paragraph_format
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)

    # 显式归零"文本之前/之后"（防止 WPS 显示错位）
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:rightChars'), '0')

    if chars == 0:
        pf.first_line_indent = Pt(0)
        ind.set(qn('w:firstLineChars'), '0')
        ind.set(qn('w:firstLine'), '0')
        return

    # 估算 1 字符 ≈ 字号（pt）。三号(16pt)下，2 字符 ≈ 32pt
    font_size = paragraph.runs[0].font.size if paragraph.runs and paragraph.runs[0].font.size else Pt(16)
    indent_pt = int(font_size.pt * chars)
    pf.first_line_indent = Pt(indent_pt)
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    ind.set(qn('w:firstLine'), str(indent_pt * 20))


def _set_indent_chars(paragraph, left_chars=None, first_line_chars=None, hanging_chars=None):
    """仅设置缩进的 Chars 字符单位属性（不触 w:left/w:hanging/w:firstLine 磅值）。

    WPS 优先用 Chars 单位显示"文本之前/之后"。磅值由 pf.left_indent 等属性单独设置。
    """
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)

    if left_chars is not None:
        ind.set(qn('w:leftChars'), str(int(left_chars * 100)))

    if first_line_chars is not None:
        ind.set(qn('w:firstLineChars'), str(int(first_line_chars * 100)))
        # 清除悬挂属性，避免冲突
        for attr in ('w:hangingChars',):
            if ind.get(qn(attr)):
                ind.attrib.pop(qn(attr), None)

    if hanging_chars is not None:
        ind.set(qn('w:hangingChars'), str(int(hanging_chars * 100)))
        # 清除首行缩进属性，避免冲突
        for attr in ('w:firstLineChars',):
            if ind.get(qn(attr)):
                ind.attrib.pop(qn(attr), None)


def remove_empty_runs(paragraph):
    """清理段落的空 run"""
    for run in paragraph.runs:
        if not run.text.strip():
            run._element.getparent().remove(run._element)


def normalize_quotes(paragraph):
    """规范化引号：状态机方式区分开/闭引号（奇数次=开，偶数次=闭）

    Runs 可能把引号对切断，因此对齐到段落级统一处理后写回。

    返回: (double_count, single_count) 实际替换的半角引号数量
    """
    full_text = ''.join(r.text for r in paragraph.runs)
    # 快检：无半角引号则跳过
    dbl_count = full_text.count('"')
    sgl_count = full_text.count("'")
    if dbl_count == 0 and sgl_count == 0:
        return (0, 0)

    # 状态机：奇数次出现=左引号，偶数次=右引号
    result = []
    in_quote = False
    for ch in full_text:
        if ch == '"':
            result.append('\u201c' if not in_quote else '\u201d')
            in_quote = not in_quote
        elif ch == "'":
            result.append('\u2018' if not in_quote else '\u2019')
            in_quote = not in_quote
        else:
            result.append(ch)
    new_text = ''.join(result)

    # 写回 runs：第一个 run 承载全部文本，其余清空
    for r in paragraph.runs:
        r.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)

    return (dbl_count, sgl_count)


def trim_paragraphs(paragraphs):
    """清理所有段落 runs 的首尾空白字符"""
    for p in paragraphs:
        for run in p.runs:
            if run.text:
                run.text = run.text.strip()


def remove_empty_paragraphs(doc):
    """删除文档 body 中的空段落，但保留附件前和文档末尾区域的空行。

    - 附件前空行保留（作为附件与正文的段间距）
    - 附件和落款日期之间保留 1 个空行
    - 文档末尾的空行保留
    """
    body = doc.element.body

    # 仅 w:p 计数（不含表格等）
    total_paragraphs = sum(1 for c in body if c.tag.endswith('}p'))

    # 第一遍：识别所有附件和日期段落位置
    attachment_p = set()  # 附件段落的段落序号
    date_p = set()       # 日期段落的段落序号
    p_index = 0
    for child in body:
        if child.tag.endswith('}p'):
            texts = []
            for r in child.iter(qn('w:t')):
                if r.text:
                    texts.append(r.text)
            full_text = ''.join(texts).strip()
            if full_text and ATTACHMENT_PATTERN.match(full_text):
                attachment_p.add(p_index)
            if full_text and DATE_PATTERN.match(full_text):
                date_p.add(p_index)
            p_index += 1

    # 第二遍：判断每个空段是否需要保留
    p_index = 0
    to_remove = []
    last_kept_attachment = -1
    for child in body:
        if child.tag.endswith('}p'):
            texts = []
            for r in child.iter(qn('w:t')):
                if r.text:
                    texts.append(r.text)
            full_text = ''.join(texts).strip()
            if not full_text:
                # 规则1：附件前的空段保留（仅保留直接相邻附件的 1 个）
                is_before_attachment = False
                nxt = _next_paragraph_elem(child, body)
                if nxt is not None:
                    nxt_texts = []
                    for r in nxt.iter(qn('w:t')):
                        if r.text:
                            nxt_texts.append(r.text)
                    nxt_full = ''.join(nxt_texts).strip()
                    if nxt_full and ATTACHMENT_PATTERN.match(nxt_full):
                        is_before_attachment = True

                # 规则2：附件块与签名块之间保留 1 个空段
                # 签名块包括：落款单位、落款日期、联系人、版记
                SIGNATURE_LABELS = (
                    'date_signature', 'unit_signature',
                    'contact', 'banji', 'signer', 'doc_number'
                )
                is_between_attach_and_signature = False
                if not is_before_attachment and p_index > max(attachment_p, default=-1):
                    nxt_p = _next_paragraph_elem(child, body)
                    if nxt_p is not None:
                        nxt_texts = []
                        for r in nxt_p.iter(qn('w:t')):
                            if r.text:
                                nxt_texts.append(r.text)
                        nxt_full = ''.join(nxt_texts).strip()
                        # 下一段是落款日期 或 短无标点（可能是落款单位）
                        if nxt_full and (
                            DATE_PATTERN.match(nxt_full) or
                            (len(nxt_full) < 30 and not nxt_full.endswith(('。', '！', '？', '.', '．', '：', ':')))
                        ):
                            is_between_attach_and_signature = True

                # 规则3：文档末尾 1 个空段保留
                is_last = p_index == total_paragraphs - 1

                keep = is_before_attachment or is_between_attach_and_signature or is_last
                if keep:
                    # 保留的空段：统一应用 28pt 行距，避免高度不一致
                    pPr = child.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        child.insert(0, pPr)
                    # 设置 spacing
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:line'), '560')  # 28pt = 560 (20ths)
                    spacing.set(qn('w:lineRule'), 'exact')
                    spacing.set(qn('w:before'), '0')
                    spacing.set(qn('w:after'), '0')
                else:
                    to_remove.append(child)
            p_index += 1
    for elem in to_remove:
        body.remove(elem)
    return len(to_remove)


def _next_paragraph_elem(p_elem, body):
    """返回 body 中 p_elem 的下一个 w:p 兄弟（跳过表格等）"""
    found = False
    for child in body:
        if child is p_elem:
            found = True
            continue
        if found and child.tag.endswith('}p'):
            return child
    return None


def normalize_bullets(paragraph):
    """规范化项目符号：替换非标符号（•·●■◆ 等）为规范格式"""
    text = paragraph.text.strip()
    if not text:
        return
    bullet_patterns = [
        (r'^[\u00b7\u2022\u25cf\u25a0\u2726\u25c6\uff0a\u2023\u2043\uff65]\s*', '\u2014\u2014'),
    ]
    for pattern, replacement in bullet_patterns:
        m = re.match(pattern, text)
        if m:
            for run in paragraph.runs:
                if run.text:
                    run.text = re.sub(pattern, replacement, run.text, count=1)
                    break
            break


def setup_page(doc):
    """设置页面边距、奇偶页页码"""
    # 启用奇偶页不同的页眉/页脚
    settings = doc.settings.element
    evenAndOdd = settings.find(qn('w:evenAndOddHeaders'))
    if evenAndOdd is None:
        evenAndOdd = OxmlElement('w:evenAndOddHeaders')
        settings.append(evenAndOdd)

    for section in doc.sections:
        section.top_margin = PAGE_TOP
        section.bottom_margin = PAGE_BOTTOM
        section.left_margin = PAGE_LEFT
        section.right_margin = PAGE_RIGHT
        section.footer_distance = FOOTER_DISTANCE

        # 奇数页页码（右下角）
        _add_page_number(section.footer, align_right=True)
        # 偶数页页码（左下角）
        _add_page_number(section.even_page_footer, align_right=False)


def _add_page_number(footer, align_right):
    """在页脚添加"— N —"格式页码（宋体四号），页脚外侧：奇右/偶左"""
    footer.is_linked_to_previous = False

    # 清空页脚原有内容
    for p in footer.paragraphs:
        p.clear()

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(28)

    # 左短横
    run1 = p.add_run("- ")
    set_font(run1, FONT_PAGE_NUM, SIZE_PAGE_NUM)

    # PAGE 域
    run2 = p.add_run()
    set_font(run2, FONT_PAGE_NUM, SIZE_PAGE_NUM)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run2._element.append(fldChar_begin)
    run2._element.append(instrText)
    run2._element.append(fldChar_end)

    # 右短横
    run3 = p.add_run(" -")
    set_font(run3, FONT_PAGE_NUM, SIZE_PAGE_NUM)


def format_all_runs(paragraph, font_name, font_size, bold=False):
    """格式化段落内所有 run"""
    for run in paragraph.runs:
        set_font(run, font_name, font_size, bold)


def body_paragraphs(doc):
    """返回文档顶层段落（排除表格单元格内的段落），用于正文分类与排版。
    避免表格单元格里的"1.""序号"等内容被误判为正文层级或覆盖表格字体。"""
    return [p for p in doc.paragraphs if p._p.getparent().tag != qn('w:tc')]


def split_heading_paragraph(paragraph, heading_label):
    """将"（一）xxx：正文"或"1、标题：正文"格式的标题段落拆成两段。

    拆分点优先级：冒号 > 破折号 > 句号（仅当句号后还有正文时使用）
    拆分后：
      - 原段落只保留标题部分（如"（一）私有化企业微信"）
      - 紧随其后插入一个新段落，包含原标题后的正文

    返回插入的新段落（如果有），否则返回 None。
    """
    import copy

    text = paragraph.text
    if not text:
        return None

    # 找拆分点：依次查找 ：、——、。（冒号优先；句号仅在正文非空时使用，避免吃掉末尾句号）
    split_marker = None
    split_pos = -1
    for marker in ['\uff1a', '\u2014\u2014', '\u3002']:  # ：、——、。
        pos = text.find(marker)
        if pos == -1 or pos < 4:
            continue
        # 关键修复：确保正文部分非空，避免被末尾的句号拦截
        body_test = text[pos + len(marker):].strip()
        if not body_test:
            continue
        split_marker = marker
        split_pos = pos
        break

    if split_marker is None:
        return None

    head_text = text[:split_pos].rstrip()
    body_text = text[split_pos + len(split_marker):].strip()

    if not body_text:
        return None

    # 正文过短（如"——制度咨询智能体"仅为标题后缀），不拆分
    MIN_BODY_LENGTH = 15  # 正文至少 15 字才独立成段，否则视为标题副标题
    if len(body_text) < MIN_BODY_LENGTH:
        return None

    # 复制原段落元素并插入到其后
    p_elem = paragraph._element
    new_p_elem = copy.deepcopy(p_elem)

    # 处理 new_p_elem：清除其所有 run，将正文写入第一个 run
    new_runs = new_p_elem.findall(qn('w:r'))
    for r in new_runs:
        r.getparent().remove(r)

    new_run = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.set(qn('xml:space'), 'preserve')
    new_t.text = body_text
    new_run.append(new_t)
    pPr = new_p_elem.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(new_run)
    else:
        new_p_elem.insert(0, new_run)

    # 处理原段落：清除其所有 run，将标题文本写入
    old_runs = p_elem.findall(qn('w:r'))
    for r in old_runs:
        r.getparent().remove(r)

    old_run = OxmlElement('w:r')
    old_t = OxmlElement('w:t')
    old_t.set(qn('xml:space'), 'preserve')
    old_t.text = head_text
    old_run.append(old_t)
    pPr_old = p_elem.find(qn('w:pPr'))
    if pPr_old is not None:
        pPr_old.addnext(old_run)
    else:
        p_elem.insert(0, old_run)

    # 插入新段落
    p_elem.addnext(new_p_elem)
    return new_p_elem


def apply_body_format(paragraph, first_indent=True):
    """对段落应用正文格式"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
    pf.left_indent = Pt(0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BODY, SIZE_BODY)
    if first_indent:
        set_first_line_indent_chars(paragraph, 2)
    else:
        set_first_line_indent_chars(paragraph, 0)


def normalize_attachment_text(paragraph):
    """清理附件说明文本：
    - 去除首尾的《》书名号
    - 去除文件名末尾的中英文标点（。，．,;；）
    """
    for run in paragraph.runs:
        if not run.text:
            continue
        text = run.text
        text = re.sub(r'^《', '', text)
        text = re.sub(r'》$', '', text)
        text = re.sub(r'[。，．,;；]+$', '', text)
        if text != run.text:
            run.text = text


# 半角→全角：仅在标题/附件场景下调用
HALF_TO_FULL_DIGIT_DOT = re.compile(r'(\d+)\.(?!\d)')


def normalize_serial_dots(paragraph):
    """将标题/附件序号中的半角句点替换为全角句点（'1.' → '1．'）。
    仅在 heading/attachment 段落调用；正文不动。"""
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = HALF_TO_FULL_DIGIT_DOT.sub(r'\1．', run.text)
        if new_text != run.text:
            run.text = new_text


# ── 段落分类 ──────────────────────────────────────────────

# 标题关键词匹配（必须以"关于"开头，避免正文中引用文件名被误判）
TITLE_PATTERNS = [
    re.compile(r'^关于.+的(报告|请示|函|通知|通报|批复|决定|意见|纪要|议案)\s*$'),
]

# 主送机关
RECIPIENT_PATTERN = re.compile(r'^[^：:]+[：:]$')

# 层级标题
HEADING_1_PATTERN = re.compile(r'^[一二三四五六七八九十]+、')         # 一、
HEADING_2_PATTERN = re.compile(r'^（[一二三四五六七八九十]+）')        # （一）
HEADING_3_PATTERN = re.compile(r'^\d+[\.\、\．]')                       # 1. 或 1、或 1．
HEADING_4_PATTERN = re.compile(r'^（\d+）')                           # （1）

# 图片标题（架构图/流程图/示意图等）：居中、仿宋小四、无首行缩进
IMAGE_CAPTION_PATTERN = re.compile(r'.*(架构图|流程图|示意图|效果图|结构图|框架图|拓扑图|部署图|界面图|样图|截图)\s*$')

# —— 开头段落（疑似未编号小标题，待用户确认是否转为正式标题）
# 同时匹配 normalize_bullets 会转换的各种 bullet 字符（·•●■◆✦＊‣⁃･）
BULLET_HEADING_PATTERN = re.compile(r'^[\u2014\u2014\u00b7\u2022\u25cf\u25a0\u2726\u25c6\uff0a\u2023\u2043\uff65\u2024]\s*\S+')

# 附件
ATTACHMENT_PATTERN = re.compile(r'^附件[：:\s]*\d*')

# 落款（日期）— 支持 "2024年6月XX日" 等占位格式
DATE_PATTERN = re.compile(r'^\d{4}年\d{1,2}月(\d{1,2}|[Xx]+)日$')

# 联系人
CONTACT_PATTERN = re.compile(r'[（(]\s*联系人\s*[：:].+[）)]?$')

# 发文字号（如 深港集团发〔2026〕5号）
DOC_NUMBER_PATTERN = re.compile(r'〔\d{4}〕.*号$')
# 签发人（上行文：报告/请示）
SIGNER_PATTERN = re.compile(r'签发人[：:]')
# 版记：抄送机关 / 印发机关及日期
BANJI_PATTERN = re.compile(r'^(抄送[：:]?|主送[：:]?)|印发')

# 文种结尾
REPORT_END = re.compile(r'(专此报告|特此报告)[。.]?$')
QINGSHI_END = re.compile(r'(妥否，请批复|专此请示|请批复)[。.]?$')
HAN_END = re.compile(r'(特此函达|请函复|此函|此复)[。.]?$')


def classify_paragraph(text, index, total):
    """根据文本内容和位置分类段落"""
    text_stripped = text.strip()
    if not text_stripped:
        return 'empty'

    # 文档首段非空 → 标题候选
    if index == 0:
        return 'title_candidate'

    # 层级标题 — 必须优先于标题关键词，避免正文中的"关于"等词误触发
    if HEADING_1_PATTERN.match(text_stripped):
        return 'heading1'
    if HEADING_2_PATTERN.match(text_stripped):
        return 'heading2'
    if HEADING_3_PATTERN.match(text_stripped):
        return 'heading3'
    if HEADING_4_PATTERN.match(text_stripped):
        return 'heading4'

    # —— 开头段落（疑似未编号小标题，待用户确认转正式标题）
    if BULLET_HEADING_PATTERN.match(text_stripped) and len(text_stripped) > 4:
        return 'bullet_heading'

    # 标题关键词
    for pat in TITLE_PATTERNS:
        if pat.search(text_stripped):
            return 'title'

    # 主送机关（文档前部，以"："结尾，短行，不含层级标题特征）
    if (index > 0 and index < max(3, total // 5)
            and RECIPIENT_PATTERN.match(text_stripped)
            and len(text_stripped) < 25):  # 真正的主送机关通常很短
        return 'recipient'

    # 附件
    if ATTACHMENT_PATTERN.match(text_stripped):
        return 'attachment'

    # 联系人
    if CONTACT_PATTERN.search(text_stripped):
        return 'contact'

    # 发文字号
    if DOC_NUMBER_PATTERN.search(text_stripped):
        return 'doc_number'

    # 签发人
    if SIGNER_PATTERN.search(text_stripped):
        return 'signer'

    # 版记（抄送/印发）
    if BANJI_PATTERN.search(text_stripped):
        return 'banji'

    # 图片标题（如"XX架构图"）：居中、仿宋小四
    if IMAGE_CAPTION_PATTERN.match(text_stripped):
        return 'image_caption'

    # 日期落款（文档末尾附近）
    if index > total * 0.7 and DATE_PATTERN.match(text_stripped):
        return 'date_signature'

    # 文种结尾
    if REPORT_END.search(text_stripped) or QINGSHI_END.search(text_stripped) or HAN_END.search(text_stripped):
        return 'end_text'

    return 'body'


# ── 格式化应用 ────────────────────────────────────────────


TITLE_LINE_BREAK_THRESHOLD = 18  # 标题超过此字数时尝试断行

# 标题断行候选词组合：双字动词/动宾短语（如 打造、推动、构建、助力、获国家 等）
# 第一组：常见动词首字
TITLE_BREAK_CHAR1 = (
    '\u643a\u6253\u63a8\u6784\u52a9\u4ee5\u4e3a\u5171\u540c\u8054\u5408'
    '\u5f00\u542f\u6df1\u9a71\u83b7\u8363\u65a9\u593a\u53d6\u6458'
)  # 携打推构助以为共同联合开启深驱获荣斩夺取摘
# 第二组：动词尾字 / 常见宾语首字
TITLE_BREAK_CHAR2 = (
    '\u9009\u5b9a\u5165\u7b51\u8d4b\u63d0\u6df1\u9a71\u878d\u805a'
    '\u9020\u52a8\u529b\u80fd\u5347\u56fd\u5e02\u7814\u521b\u53d1\u5c55'
)  # 选定入筑赋提深驱融聚造动力能升国市研创发展

TITLE_BREAK_POINTS = re.compile(
    f'[{TITLE_BREAK_CHAR1}][{TITLE_BREAK_CHAR2}]'
)

# 标题断行保护词：断点落在此类固定名词内部时跳过该候选
# 每项为 (prefix, suffix) 形式——prefix 为词的前半部分，检查断点是否在词内部
UNBREAKABLE_PHRASES = [
    '国家发展', '国家发改', '国家发改委', '国家发展和改革委员会',
    '深圳市', '广东省', '北京市', '上海市', '广州市', '天津市', '重庆市',
    '人民政府', '国务院', '党中央', '中央军委',
    '研究成果', '深圳港', '中国', '集团',
    '高质量', '现代化', '信息化', '数字化',
    '社会主义', '中国特色', '中华民族',
    '海洋产业', '海洋强国', '全球海洋',
    '发改委', '交通运输', '科技创新',
]


def _is_break_safe(text, pos):
    """检查在 pos 处断行是否安全（不会割裂固定名词）

    检查 pos 前后共 3-6 个字的窗口是否命中不可断名词列表。
    """
    # 断点前的词尾（最多取3字）
    prefix = text[max(0, pos - 3):pos]
    # 断点后的词头（最多取4字）
    suffix = text[pos:min(len(text), pos + 4)]

    for phrase in UNBREAKABLE_PHRASES:
        plen = len(phrase)
        # 遍历所有可能的切分点：phrase[:k] 属前、phrase[k:] 属后
        for k in range(1, plen):
            if phrase[:k] == prefix[-k:] and phrase[k:] == suffix[:plen - k]:
                return False
    return True


def _find_title_break_options(text):
    """返回标题所有候选断行位置，每个候选为 (pos, label, preview_text)

    按优先级排序，最多 3 个候选。preview_text 用换行符示意。
    断行位置自动避开固定名词（如"国家发改委"、"深圳市"等）。
    """
    if len(text) <= TITLE_LINE_BREAK_THRESHOLD:
        return []
    options = []

    # 候选 A：语义断点（动词组合前）
    for m in TITLE_BREAK_POINTS.finditer(text):
        pos = m.start()
        if 4 < pos < len(text) - 4:  # 不在两端
            if not _is_break_safe(text, pos):
                continue
            preview = text[:pos] + '\u21b5' + text[pos:]
            options.append((pos, '\u8bed\u4e49\u65ad\u70b9', preview))
            if len(options) >= 1:
                break

    # 候选 B：2/3 处标点/空格前
    mid = len(text) * 2 // 3
    for i in range(mid, min(len(text) - 2, mid + 6)):
        if text[i] in ('\uff0c', '\u3001', ' ', '\u3000'):
            p = i + 1
            if not _is_break_safe(text, p):
                continue
            preview = text[:p] + '\u21b5' + text[p:]
            options.append((p, '\u6807\u70b9\u65ad\u70b9', preview))
            break

    # 候选 C：1/2 处（居中）—— 仅在安全时才加入
    half = len(text) // 2
    if _is_break_safe(text, half):
        preview = text[:half] + '\u21b5' + text[half:]
        options.append((half, '\u5c45\u4e2d\u65ad\u70b9', preview))

    # 去重
    seen = set()
    unique = []
    for pos, label, preview in options:
        if pos not in seen:
            seen.add(pos)
            unique.append((pos, label, preview))
    return unique[:3]


def apply_title_format(paragraph, break_pos=None):
    """标题格式：方正小标宋二号居中，可选断行

    Args:
        break_pos: 如果指定，在此位置插入换行符；None 则不断行
    """
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_TITLE, SIZE_TITLE)

    if break_pos is None:
        return

    import copy

    full_text = paragraph.text
    if break_pos <= 0 or break_pos >= len(full_text):
        return

    runs_data = [(r, r.text) for r in paragraph.runs]
    if not runs_data:
        return

    # 定位断行点落在哪个 run、以及该 run 内的偏移
    char_count = 0
    break_idx = None
    offset = 0
    for idx, (r, t) in enumerate(runs_data):
        run_len = len(t)
        if char_count + run_len > break_pos:
            break_idx = idx
            offset = break_pos - char_count
            break
        char_count += run_len
    if break_idx is None:
        return

    # 先清空所有 run 的文本（保留 run 节点与格式属性）
    for r, _ in runs_data:
        r.text = ''

    # 还原断行点之前的 run 文本（这些 run 的字符全部在断行点之前）
    for idx in range(break_idx):
        runs_data[idx][0].text = runs_data[idx][1]

    # 拆分断行点所在 run：左侧保留，右侧文本（含其后所有 run）放到新 run
    r, t = runs_data[break_idx]
    r.text = t[:offset]
    after_part = full_text[break_pos:]
    if after_part:
        br = OxmlElement('w:br')
        r._element.append(br)
        new_run = OxmlElement('w:r')
        rpr = r._element.find(qn('w:rPr'))
        if rpr is not None:
            new_run.append(copy.deepcopy(rpr))
        new_t = OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = after_part
        new_run.append(new_t)
        r._element.addnext(new_run)


def apply_recipient_format(paragraph):
    """主送机关格式：仿宋三号顶格"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BODY, SIZE_BODY)


def apply_heading1_format(paragraph):
    """一级标题：黑体三号"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 2)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_HEITI, SIZE_BODY)


def apply_heading2_format(paragraph):
    """二级标题：楷体国标三号加粗"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 2)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_KAITI, SIZE_BODY, bold=True)


def apply_heading3_format(paragraph):
    """三级标题：仿宋三号"""
    apply_body_format(paragraph)
    # 三级标题号后不加粗
    normalize_serial_dots(paragraph)


def apply_heading4_format(paragraph):
    """四级标题：仿宋三号"""
    apply_body_format(paragraph)
    normalize_serial_dots(paragraph)


def apply_attachment_format(paragraph, attachment_index=1, total_attachments=1):
    """附件格式：仿宋三号，按附件序号和总数应用不同的悬挂缩进。

    规则（参考用户实际公文格式）：
    - 1个附件（无序号）: 左缩进0，首行缩进2字符
    - 附件1: 左缩进6.5字符，首行缩进-4.5字符（"1."前2字符起始，续行6.5字符）
    - 附件2+: 左缩进6.5字符，首行缩进-1.5字符（"N."前5字符起始，续行6.5字符）

    关键：同时设置 w:left/leftChars（磅+字符双单位），避免源文档遗留的
    leftChars=0 等属性覆盖新的 left 值（WPS 优先用 Chars 单位）。
    """
    # 文本规范化：去书名号、去末尾标点
    normalize_attachment_text(paragraph)
    # 半角→全角序号点（仅在附件场景）
    normalize_serial_dots(paragraph)

    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    char_pt = SIZE_BODY.pt  # 三号字单字符宽度

    if total_attachments == 1:
        # 1个附件（无序号）：左缩进0，首行缩进2字符
        pf.left_indent = Pt(0)
        pf.first_line_indent = Pt(2 * char_pt)
        _set_indent_chars(paragraph, left_chars=0, first_line_chars=2)
    elif attachment_index == 1:
        # 附件1：left=6.5ch, hanging=4.5ch
        # leftChars/hangingChars 按参考文档值（WPS 优先用 Chars 单位显示"文本之前"）
        left = 6.5 * char_pt
        hang = 4.5 * char_pt
        pf.left_indent = Pt(left)
        pf.first_line_indent = Pt(-hang)
        _set_indent_chars(paragraph, left_chars=3.05, hanging_chars=4.5)
    else:
        # 附件2+：left=6.5ch, hanging=1.5ch
        # leftChars/hangingChars 按参考文档值
        left = 6.5 * char_pt
        hang = 1.5 * char_pt
        pf.left_indent = Pt(left)
        pf.first_line_indent = Pt(-hang)
        _set_indent_chars(paragraph, left_chars=7.62, hanging_chars=1.5)

    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BODY, SIZE_BODY)


def apply_signature_format(paragraph):
    """落款格式：仿宋三号右对齐"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BODY, SIZE_BODY)


def apply_signature_unit_format(paragraph):
    """落款单位格式：仿宋三号右对齐"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BODY, SIZE_BODY)


def apply_contact_format(paragraph):
    """联系人格式：仿宋三号"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 2)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BODY, SIZE_BODY)


def apply_end_text_format(paragraph):
    """文种结尾用语"""
    apply_body_format(paragraph)


def apply_table_format(table):
    """表格格式：仿宋小四，表头行加粗，细单线边框（六向），单元格左对齐。
    列宽与已有单元格水平对齐方式保持不变（仅补全缺失边框、统一字体与表头加粗）。"""
    tblPr = table._tbl.tblPr

    # 边框：缺失则补全为细单线（六向 single / sz=4 / 黑色）
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = borders.find(qn('w:' + edge))
        if e is None:
            e = OxmlElement('w:' + edge)
            borders.append(e)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(TABLE_BORDER_SZ))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), TABLE_BORDER_COLOR)

    # 表格整体对齐：已有则保留，缺失默认左对齐（与正文对齐）
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        tblPr.append(jc)
        jc.set(qn('w:val'), 'left')

    for ri, row in enumerate(table.rows):
        is_header = (ri == 0)
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # 单元格水平对齐：已有则保留，缺失默认左对齐
            cjc = tcPr.find(qn('w:jc'))
            if cjc is None:
                cjc = OxmlElement('w:jc')
                tcPr.append(cjc)
                cjc.set(qn('w:val'), 'left')
            # 垂直对齐：缺失则居中（更美观）
            vjc = tcPr.find(qn('w:vAlign'))
            if vjc is None:
                vjc = OxmlElement('w:vAlign')
                tcPr.append(vjc)
                vjc.set(qn('w:val'), 'center')
            bold = TABLE_HEADER_BOLD if is_header else False
            for p in cell.paragraphs:
                set_line_spacing(p, LINE_SPACING)
                format_all_runs(p, FONT_TABLE, SIZE_TABLE, bold=bold)


def apply_doc_number_format(paragraph):
    """发文字号：仿宋三号居中"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_DOCNUM, SIZE_DOCNUM)


def apply_signer_format(paragraph):
    """签发人：仿宋三号右对齐（上行文与发文字号同行）"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_SIGNER, SIZE_SIGNER)


def apply_banji_format(paragraph):
    """版记（抄送/印发）：仿宋四号"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BANJI, SIZE_BANJI)


def apply_image_caption_format(paragraph):
    """图片标题（如"XX架构图"）：仿宋小四、居中、无首行缩进"""
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.left_indent = Pt(0)
    set_first_line_indent_chars(paragraph, 0)
    set_line_spacing(paragraph, LINE_SPACING)
    format_all_runs(paragraph, FONT_BODY, SIZE_TABLE)  # 仿宋小四(14pt)


# ── 后处理：智能识别落款 ──────────────────────────────────


def _fix_sequential_attachments(paragraphs, labels):
    """修复连续附件编号：将"附件：1.xxx"后的"2.xxx""3.xxx"重标为附件。

    源文档中常见格式：第一行写"附件：1.xxx"，后续行写"2.xxx""3.xxx"
    但没带"附件"前缀，被 classify_paragraph 误判为 body/heading3。
    此函数扫描并纠正，不添加前缀（参考文档中附件2/3就是纯数字编号）。
    """
    texts = [p.text.strip() for p in paragraphs]
    for i, label in enumerate(labels):
        if label == 'attachment':
            attach_text = texts[i]
            attach_match = re.match(r'^附件[：:\s]*(\d+)[\.\．]', attach_text)
            if attach_match:
                base_num = int(attach_match.group(1))
                next_num = base_num + 1
                for j in range(i + 1, len(labels)):
                    if labels[j] in ('heading3', 'heading4', 'body'):
                        next_text = texts[j]
                        if re.match(rf'^{next_num}[\.\．]', next_text):
                            labels[j] = 'attachment'
                            next_num += 1
                        else:
                            break
                    elif labels[j] in ('heading1', 'heading2', 'attachment', 'empty'):
                        continue
                    else:
                        break


def _fix_unnumbered_headings(paragraphs, labels):
    """将短正文段落提升为 heading1（无编号标题）。

    两种模式：
    1. 热启动：文档已有 heading1，从其位置向后扫描提升后续短正文
    2. 冷启动：文档无任何 heading1，按段间特征检测一级标题（短正文+后有更长正文）
    """
    texts = [p.text.strip() for p in paragraphs]
    promoted = 0
    has_heading1 = any(lb == 'heading1' for lb in labels)

    if has_heading1:
        # 热启动：从已有 heading1 向后扫描
        last_heading1 = -1
        for i in range(len(labels)):
            if labels[i] == 'heading1':
                last_heading1 = i
            elif labels[i] == 'body' and last_heading1 >= 0:
                t = texts[i]
                if len(t) < 30 and not t.endswith(('。', '！', '？', '.', '．')):
                    if i + 1 < len(labels) and labels[i + 1] == 'body':
                        if len(texts[i + 1]) >= len(t):
                            labels[i] = 'heading1'
                            last_heading1 = i
                            promoted += 1
    else:
        # 冷启动：扫描所有短正文作为候选一级标题（<20字、无句末标点、后有更长段落）
        # 过滤规则：不以"主要"开头、不以"："结尾（通常是列表引导句，非标题）
        last_heading1 = -1
        for i in range(len(labels)):
            if labels[i] == 'body':
                t = texts[i]
                # 过滤：以"主要"开头的子标题、以"："结尾的引导句
                if t.startswith('主要') or t.endswith('：') or t.endswith(':'):
                    continue
                if len(t) < 20 and not t.endswith(('。', '！', '？', '.', '．')):
                    if i + 1 < len(labels) and labels[i + 1] in ('body', 'heading2', 'heading3', 'heading4'):
                        if texts[i + 1] and len(texts[i + 1]) >= len(t):
                            labels[i] = 'heading1'
                            last_heading1 = i
                            promoted += 1
    return promoted


def _ensure_title_paragraph(labels, texts):
    """确认标题段落（支持首段空段场景）并合并多段标题。

    3 步：
    1. 如果有 'title'：保留
    2. 如果有 'title_candidate'：提升为 'title'
    3. 如果首段是 empty：找下一个非空短正文（<30字、无句末标点）作为 title
    4. 标题提升后，紧随的短正文也并入 title（处理手工换行标题）
    """
    title_found = False
    title_idx = -1
    for i, label in enumerate(labels):
        if label == 'title':
            title_found = True
            title_idx = i
            break
    if not title_found:
        for i, label in enumerate(labels):
            if label == 'title_candidate':
                labels[i] = 'title'
                title_found = True
                title_idx = i
                break
    if not title_found:
        skip = ['集团领导：', '集团：', '各部门：', '各有关单位：', '尊敬的', '您好']
        for i, label in enumerate(labels):
            if label == 'body' and texts[i].strip():
                t = texts[i].strip()
                if len(t) < 30 and not t.endswith(('。', '！', '？', '.', '．', '：', ':')):
                    if not any(t.startswith(s) for s in skip):
                        labels[i] = 'title'
                        title_found = True
                        title_idx = i
                        break

    # 多段标题合并
    if title_idx >= 0:
        for j in range(title_idx + 1, len(labels)):
            lb = labels[j]
            if lb in ('heading1', 'heading2', 'heading3', 'heading4', 'recipient',
                      'attachment', 'date_signature', 'unit_signature', 'contact',
                      'doc_number', 'signer', 'banji', 'end_text', 'image_caption'):
                break
            if lb == 'body':
                t = texts[j].strip()
                if t and len(t) < 30 and not t.endswith(('。', '！', '？', '.', '．', '：', ':')):
                    skip2 = ['集团领导：', '集团：', '各部门：', '各有关单位：']
                    if not any(t.startswith(s) for s in skip2):
                        labels[j] = 'title'
                        continue
                break
            else:
                break


def post_process_signatures(paragraphs, labels):
    """在文档末尾智能识别落款单位和日期"""
    total = len(paragraphs)

    # 从文档末尾往前查找可能的单位名称和日期
    guess_company = None
    guess_date = None

    for i in range(total - 1, max(0, total - 15), -1):
        text = paragraphs[i].text.strip()
        if not text:
            # 空段落 → 可能是间距
            continue

        # 保护：附件段落不参与落款识别
        if labels[i] == 'attachment':
            continue

        # 日期段落
        if DATE_PATTERN.match(text):
            if labels[i] != 'date_signature':
                labels[i] = 'date_signature'
            guess_date = i
            continue

        # 单位名称段落（在日期上方，长度适中，不含标点）
        if guess_date and i < guess_date:
            if len(text) > 4 and len(text) < 30 and not re.search(r'[，。！？、：；""''（）《》]', text):
                if not any(pat.search(text) for pat in TITLE_PATTERNS):
                    labels[i] = 'unit_signature'
                    guess_company = i
                    break


# ── 结构预览 ──────────────────────────────────────────────

LABEL_NAMES = {
    'title': '【标题】', 'title_candidate': '【标题候选】',
    'recipient': '【主送机关】', 'heading1': '一、', 'heading2': '（一）',
    'heading3': '1.', 'heading4': '（1）', 'attachment': '【附件】',
    'date_signature': '【落款日期】', 'unit_signature': '【落款单位】',
    'contact': '【联系人】', 'end_text': '【结尾】', 'body': '【正文】',
    'bullet_heading': '【——小标题】',
    'doc_number': '【发文字号】', 'signer': '【签发人】', 'banji': '【版记】',
    'image_caption': '【图片标题】',
    'empty': '',
}


def detect_hierarchy_issues(labels, texts):
    """检测层级跳跃问题（如 一、下直接出现 1、而没有（一））"""
    issues = []
    for i, label in enumerate(labels):
        if label == 'heading3' and i > 0:
            for j in range(i - 1, -1, -1):
                prev = labels[j]
                if prev == 'heading1':
                    issues.append((i, texts[i].strip()[:50], '\u26a0 "一、"下直接出现"1、"（建议改为"（一）"）'))
                    break
                elif prev == 'heading2':
                    break
                elif prev == 'heading3':
                    continue
                elif prev in ('body', 'empty'):
                    continue
                else:
                    break
    return issues


def diagnose_document(doc, paragraphs, labels):
    """全要素诊断：返回排版规范符合度报告

    返回:
        dict: {
            'quotes': (count, msg),
            'page_margins': (ok, msg),
            'page_number': (ok, msg),
            'line_spacing': (ok, msg),
            'bullets': (count, msg),
            'title_break': (has_long_title, title_text, options),
        }
    """
    result = {}

    # --- 引号统计 ---
    quote_count = 0
    for p in paragraphs:
        text = p.text
        quote_count += text.count('"') + text.count("'")
    result['quotes'] = (quote_count, f'全文 {quote_count} 处半角引号 → 自动替换为全角配对接号')

    # --- 页边距 ---
    try:
        sect = doc.sections[0]
        margins_ok = True
        margin_details = []
        checks = [
            ('上', sect.top_margin, PAGE_TOP, '3.7'),
            ('下', sect.bottom_margin, PAGE_BOTTOM, '3.5'),
            ('左', sect.left_margin, PAGE_LEFT, '2.8'),
            ('右', sect.right_margin, PAGE_RIGHT, '2.6'),
        ]
        for name, actual, expected, expected_str in checks:
            if actual is not None and abs(actual - expected) > Cm(0.2):
                margins_ok = False
                margin_details.append(f'{name}边距 {actual.cm:.1f}cm（标准 {expected_str}cm）')
        if margins_ok:
            result['page_margins'] = (True, '页边距符合标准')
        else:
            result['page_margins'] = (False, '；'.join(margin_details) + ' → 自动修正为标准值')
    except Exception:
        result['page_margins'] = (True, '页边距将自动设置为标准值')

    # --- 页码 ---
    try:
        footer = doc.sections[0].footer
        has_page_field = False
        for p in footer.paragraphs:
            for run in p.runs:
                if run._element.findall(qn('w:fldChar')):
                    has_page_field = True
                    break
        if has_page_field:
            result['page_number'] = (True, '已有页码域（保留）')
        else:
            result['page_number'] = (False, '无页码 → 自动添加 "- 1 -" 格式页码（宋体四号，奇右偶左）')
    except Exception:
        result['page_number'] = (True, '页码将自动设置')

    # --- 行距 ---
    sample = None
    for p in paragraphs:
        if p.text.strip() and labels[paragraphs.index(p)] not in ('empty',):
            pf = p.paragraph_format
            if pf.line_spacing is not None:
                sample = pf.line_spacing
                break
    if sample is not None:
        try:
            sample_pt = sample.pt if hasattr(sample, 'pt') else float(sample)
            if hasattr(sample, 'pt') and abs(sample_pt - 28) < 2:
                result['line_spacing'] = (True, f'行距 {sample_pt:.0f}pt（标准 28pt 固定值）')
            else:
                result['line_spacing'] = (False, f'当前行距 ~{sample_pt:.0f}pt → 自动修正为 28pt 固定值')
        except Exception:
            result['line_spacing'] = (False, f'当前行距未标准化 → 自动修正为 28pt 固定值')
    else:
        result['line_spacing'] = (True, '行距将自动设置为 28pt 固定值')

    # --- 项目符号 ---
    bullet_count = 0
    bullet_pattern = re.compile(r'^[\u00b7\u2022\u25cf\u25a0\u2726\u25c6\uff0a\u2023\u2043\uff65\u2024]')
    for p in paragraphs:
        if bullet_pattern.match(p.text.strip()):
            bullet_count += 1
    result['bullets'] = (bullet_count, f'全文 {bullet_count} 处非标项目符号开头段落 → 将作为疑似小标题提示确认' if bullet_count else '无异常项目符号')

    # --- 长标题 ---
    title_text = ''
    for i, label in enumerate(labels):
        if label in ('title', 'title_candidate'):
            title_text = paragraphs[i].text.strip()
            break
    if len(title_text) > TITLE_LINE_BREAK_THRESHOLD:
        options = _find_title_break_options(title_text)
        result['title_break'] = (True, title_text, options)
    else:
        result['title_break'] = (False, title_text, [])

    # --- —— 开头段落（疑似未编号小标题）---
    bullet_heading_list = []
    for i, label in enumerate(labels):
        if label == 'bullet_heading':
            text = paragraphs[i].text.strip()
            bullet_heading_list.append((i, text[:60]))
    result['bullet_headings'] = bullet_heading_list
    result['table_count'] = len(doc.tables)

    return result


def fix_hierarchy_in_document(paragraphs, labels, texts):
    """在段落对象上修正层级跳跃编号（仅修改文本，不保存文件）

    策略：在每个 heading1 组内，检查 heading3 是否跳过了 heading2；
    如跳过，则将 heading3 的编号替换为 heading2 格式（（一）（二）...）。

    此函数仅修改段落文本内容，不涉及文件 I/O。

    返回: [(段号, 旧编号, 新编号), ...] 修正记录列表
    """
    LEVEL2_CHARS = "一二三四五六七八九十"

    current_h1_idx = None
    h3_counter = 0
    fixes = []

    for i, label in enumerate(labels):
        if label == 'heading1':
            current_h1_idx = i
            h3_counter = 0
        elif label == 'heading3' and current_h1_idx is not None:
            # 检查从 h1 到当前段落之间是否有 heading2
            has_heading2 = any(labels[j] == 'heading2' for j in range(current_h1_idx + 1, i))
            if not has_heading2:
                h3_counter += 1
                if h3_counter <= len(LEVEL2_CHARS):
                    new_prefix = f"（{LEVEL2_CHARS[h3_counter - 1]}）"
                    para = paragraphs[i]
                    old_text = para.text
                    new_text = re.sub(r'^\d+[.、]', new_prefix, old_text)
                    if new_text != old_text:
                        # 提取旧编号和新编号用于记录
                        old_prefix_m = re.match(r'^(\d+[.、])', old_text)
                        old_prefix = old_prefix_m.group(1) if old_prefix_m else old_text[:4]
                        para.clear()
                        para.add_run(new_text)
                        fixes.append((i, old_prefix, new_prefix))

    return fixes


def print_structure(paragraphs, labels, diagnosis=None):
    """打印文档结构预览 + 全要素诊断报告"""
    lines = []
    lines.append("\n" + "=" * 60)
    lines.append("\U0001f4cb 文档综合诊断报告")
    lines.append("=" * 60)

    # ── 1. 结构概览 ──
    cats = {'title': 0, 'heading1': 0, 'heading2': 0, 'heading3': 0,
            'heading4': 0, 'body': 0, 'attachment': 0}
    for lb in labels:
        if lb in cats:
            cats[lb] += 1
    lines.append("\n【结构概览】")
    parts = []
    if cats['title']:
        parts.append(f"标题 {cats['title']} 段")
    parts.append(f"一级标题 {cats['heading1']} 个")
    if cats['heading2']:
        parts.append(f"二级标题 {cats['heading2']} 个")
    if cats['heading3']:
        parts.append(f"三级标题 {cats['heading3']} 个")
    if cats['heading4']:
        parts.append(f"四级标题 {cats['heading4']} 个")
    parts.append(f"正文 {cats['body']} 段")
    if cats['attachment']:
        parts.append(f"附件 {cats['attachment']} 个")
    if diagnosis and diagnosis.get('table_count'):
        parts.append(f"表格 {diagnosis['table_count']} 个")
    lines.append("  " + " / ".join(parts))

    # ── 2. 段落明细 ──
    texts = [p.text.strip() for p in paragraphs]
    lines.append("\n【段落明细】")
    for i, (label, text) in enumerate(zip(labels, texts)):
        if label == 'empty' or not text:
            continue
        name = LABEL_NAMES.get(label, label)
        display = text[:60] + ('...' if len(text) > 60 else '')
        lines.append(f"  段{i:3d}  {name:<10s}  {display}")

    # ── 3. 排版规范符合度 ──
    if diagnosis:
        lines.append("\n【排版规范符合度】")
        # 引号
        q_cnt, q_msg = diagnosis.get('quotes', (0, ''))
        q_icon = '\u2705' if q_cnt == 0 else '\u2699\ufe0f'
        lines.append(f"  {q_icon} 引号规范：{q_msg}")

        # 页边距
        pm_ok, pm_msg = diagnosis.get('page_margins', (True, ''))
        pm_icon = '\u2705' if pm_ok else '\u2699\ufe0f'
        lines.append(f"  {pm_icon} 页边距：{pm_msg}")

        # 页码
        pn_ok, pn_msg = diagnosis.get('page_number', (True, ''))
        pn_icon = '\u2705' if pn_ok else '\u2699\ufe0f'
        lines.append(f"  {pn_icon} 页码：{pn_msg}")

        # 行距
        ls_ok, ls_msg = diagnosis.get('line_spacing', (True, ''))
        ls_icon = '\u2705' if ls_ok else '\u2699\ufe0f'
        lines.append(f"  {ls_icon} 行距：{ls_msg}")

        # 项目符号
        b_cnt, b_msg = diagnosis.get('bullets', (0, ''))
        b_icon = '\u2705' if b_cnt == 0 else '\u2699\ufe0f'
        lines.append(f"  {b_icon} 项目符号：{b_msg}")

        # 标题断行
        tb_info = diagnosis.get('title_break', (False, '', []))
        if tb_info[0]:
            title_text = tb_info[1]
            options = tb_info[2]
            lines.append(f"\n  \U0001f4dd 长标题（{len(title_text)}字）需断行：")
            lines.append(f"     标题：{title_text}")
            if options:
                lines.append(f"     📌 候选断行位置：")
                for idx, (pos, label, preview) in enumerate(options):
                    letter = chr(65 + idx)  # A, B, C
                    lines.append(f"       {letter}. [{label}] pos={pos}: {preview}")
                    lines.append(f"          用法: --title-break={pos}")
            else:
                lines.append(f"     ⚠ 未找到合适断点，建议手动调整")

        lines.append(f"\n  \U0001f4a1 \u2699\ufe0f = 排版时将自动修正    \u2705 = 已符合标准")

    # ── 4. 层级问题 ──
    issues = detect_hierarchy_issues(labels, texts)
    if issues:
        lines.append("\n【⚠ 层级问题】（需确认）")
        for idx, snippet, msg in issues:
            lines.append(f"  段{idx}: {msg}")
            lines.append(f"     当前: {snippet}")

    # ── 5. ——开头段落（疑似未编号小标题）──
    if diagnosis:
        bh_list = diagnosis.get('bullet_headings', [])
        if bh_list:
            lines.append("\n【⚠ ——开头段落】（疑似未编号小标题，需确认是否转为正式标题）")
            for idx, snippet in bh_list:
                lines.append(f"  段{idx}: {snippet}")

    lines.append("\n" + "=" * 60)
    return "\n".join(lines)


def generate_change_summary(diagnosis, changes):
    """排版完成后生成修改摘要，告知用户具体修改了哪些内容。

    Args:
        diagnosis: diagnose_document() 返回的诊断结果
        changes: 排版过程中收集的修改记录 dict:
            - 'quotes_double': int  双引号替换处数
            - 'quotes_single': int  单引号替换处数
            - 'title_broken': (pos, label) or None  标题断行信息
            - 'headings_split': [(段号, 标题摘要), ...]
            - 'hierarchy_fixes': [(段号, 旧编号, 新编号), ...]
            - 'bullet_headings_kept': [(段号, 内容摘要), ...]
            - 'tables_formatted': int  表格处理数
            - 'margins_fixed': bool  页边距是否修正
            - 'page_number_added': bool  是否添加了页码
            - 'line_spacing_fixed': bool  行距是否修正

    Returns:
        str: 格式化的修改摘要文本
    """
    lines = []
    lines.append("")
    lines.append("=" * 60)
    lines.append("\U0001f4dd 排版修改摘要")
    lines.append("=" * 60)

    section_count = 0

    # ── 格式修正 ──
    format_items = []

    # 页边距
    pm_ok, pm_msg = diagnosis.get('page_margins', (True, ''))
    if not pm_ok:
        format_items.append(f"页边距已修正为标准值（上3.7/下3.5/左2.8/右2.6cm）")

    # 页码
    pn_ok, pn_msg = diagnosis.get('page_number', (True, ''))
    if not pn_ok:
        format_items.append('页码已添加 "- 1 -" 格式（宋体四号，奇右偶左）')

    # 行距
    ls_ok, ls_msg = diagnosis.get('line_spacing', (True, ''))
    if not ls_ok:
        format_items.append('全文行距已修正为 28pt 固定值')

    # 段前段后
    format_items.append('段前段后已归零')

    # 空段落删除
    empty_removed = changes.get('empty_removed', 0)
    if empty_removed > 0:
        format_items.append(f'已删除 {empty_removed} 个空段落（段间空行）')

    # 表格
    if changes.get('tables_formatted', 0) > 0:
        format_items.append(
            f"{changes['tables_formatted']} 个表格已统一格式（仿宋小四、表头加粗、细单线边框）"
        )

    if format_items:
        section_count += 1
        lines.append(f"\n\u3010\u683c\u5f0f\u4fee\u6b63\u3011{len(format_items)} \u9879")
        for item in format_items:
            lines.append(f"  \u2022 {item}")

    # ── 内容规范化 ──
    content_items = []

    dbl_q = changes.get('quotes_double', 0)
    sgl_q = changes.get('quotes_single', 0)
    if dbl_q > 0 or sgl_q > 0:
        parts = []
        if dbl_q > 0:
            parts.append(f'{dbl_q} \u5904\u534a\u89d2" \u2192 \u5168\u89d2\u201c\u201d')
        if sgl_q > 0:
            parts.append(f"{sgl_q} \u5904\u534a\u89d2' \u2192 \u5168\u89d2\u2018\u2019")
        content_items.append('\u5f15\u53f7\u89c4\u8303\u5316\uff1a\u5168\u6587 ' + '\u3001'.join(parts))

    title_broken = changes.get('title_broken')
    if title_broken is not None:
        pos, label = title_broken
        content_items.append(f"\u6807\u9898\u65ad\u884c\uff1a\u7b2c{pos}\u5b57\u5904\u65ad\u5f00\uff08{label}\uff09")

    if content_items:
        section_count += 1
        lines.append(f"\n\u3010\u5185\u5bb9\u89c4\u8303\u5316\u3011{len(content_items)} \u9879")
        for item in content_items:
            lines.append(f"  \u2022 {item}")

    # ── 层级修正 ──
    hierarchy_items = []

    splits = changes.get('headings_split', [])
    if splits:
        details = []
        for idx, summary in splits:
            details.append(f"\u6bb5{idx}\"{summary[:20]}\"")
        hierarchy_items.append(
            f"\u6807\u9898\u6bb5\u843d\u62c6\u5206\uff1a{len(splits)} \u5904\uff08{' / '.join(details)}\uff09"
        )

    fixes = changes.get('hierarchy_fixes', [])
    if fixes:
        details = []
        for idx, old_pfx, new_pfx in fixes:
            details.append(f"\u6bb5{idx} \"{old_pfx}\" \u2192 \"{new_pfx}\"")
        hierarchy_items.append(
            f"\u5c42\u7ea7\u7f16\u53f7\u4fee\u6b63\uff08--fix-hierarchy\uff09\uff1a{len(fixes)} \u5904\uff08{' / '.join(details)}\uff09"
        )

    if hierarchy_items:
        section_count += 1
        lines.append(f"\n\u3010\u5c42\u7ea7\u4fee\u6b63\u3011{len(hierarchy_items)} \u9879")
        for item in hierarchy_items:
            lines.append(f"  \u2022 {item}")

    # ── 未处理 ──
    kept = changes.get('bullet_headings_kept', [])
    if kept:
        section_count += 1
        lines.append(f"\n\u3010\u672a\u5904\u7406\u3011{len(kept)} \u9879")
        for idx, snippet in kept:
            lines.append(f"  \u2022 \u6bb5{idx} \"{snippet[:30]}\" \u2014\u2014 \u4fdd\u7559\u4e3a\u6b63\u6587\u683c\u5f0f\uff08\u672a\u786e\u8ba4\u8f6c\u4e3a\u6b63\u5f0f\u6807\u9898\uff09")

    if section_count == 0:
        lines.append("\n  \u2705 \u6587\u6863\u5df2\u7b26\u5408 GB/T 9704 \u6807\u51c6\uff0c\u65e0\u9700\u4fee\u6539\u3002")

    lines.append("\n" + "=" * 60)
    return "\n".join(lines)


# ── 主流程 ────────────────────────────────────────────────


def format_document(input_path, output_path=None, preview_only=False,
                    title_break_pos=None, auto_mode=False, fix_hierarchy=False):
    """主入口：读取文档，分析结构，应用格式，保存

    Args:
        preview_only: 仅打印结构预览不保存
        title_break_pos: 标题断行位置（字符索引），None 则不断行
        auto_mode: 跳过所有确认环节，全部使用默认行为
    """

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_已排版{ext}"

    # 修改追踪记录
    changes = {
        'quotes_double': 0,
        'quotes_single': 0,
        'title_broken': None,
        'headings_split': [],
        'hierarchy_fixes': [],
        'bullet_headings_kept': [],
        'tables_formatted': 0,
    }

    doc = Document(input_path)

    # 1. 设置页面
    setup_page(doc)

    paragraphs = body_paragraphs(doc)
    total = len(paragraphs)

    # 2. 预处理：清理所有段落首尾空白
    trim_paragraphs(paragraphs)

    # 2.5 删除空段落（段间距的空行）
    removed_empty = remove_empty_paragraphs(doc)
    if removed_empty > 0:
        changes['empty_removed'] = removed_empty
        paragraphs = body_paragraphs(doc)
        total = len(paragraphs)

    # 3. 收集所有段落文本并分类（在规范化前先跑诊断，确保引号等统计准确）
    texts = [p.text for p in paragraphs]
    labels = [classify_paragraph(t, i, total) for i, t in enumerate(texts)]

    # 3.5 修复无编号标题和连续附件（必须在诊断前运行，确保诊断能检测层级问题）
    if not auto_mode:
        _fix_unnumbered_headings(paragraphs, labels)
        _fix_sequential_attachments(paragraphs, labels)

    # 4. 全要素诊断
    diagnosis = diagnose_document(doc, paragraphs, labels)

    # 5. 预处理：规范化所有段落（引号自动配对，项目符号保留原样由 bullet_heading 检测处理）
    for p in paragraphs:
        dbl, sgl = normalize_quotes(p)
        changes['quotes_double'] += dbl
        changes['quotes_single'] += sgl

    # 5.5 修复：连续附件编号（如"附件：1.xxx"后"2.xxx""3.xxx"未带"附件"前缀）
    # 将紧跟附件段落后的连续数字标题（heading3）重标为 attachment
    if not auto_mode:
        _fix_sequential_attachments(paragraphs, labels)

    # 5.6 修复：无编号一级标题（如"通过公司对公账户..."跟在 heading1 后但没写"二、"）
    _fix_unnumbered_headings(paragraphs, labels)
    _fix_sequential_attachments(paragraphs, labels)

    # 6. 后处理：确认标题段落（包含多段标题合并逻辑）
    _ensure_title_paragraph(labels, texts)

    # 8. 打印结构预览（含全要素诊断报告）
    structure_preview = print_structure(paragraphs, labels, diagnosis)

    # 仅预览模式：到此为止
    if preview_only:
        preview_msg = f"{structure_preview}\n💡 预览模式：未生成排版文件。确认结构无误后，去掉 --preview 参数重新运行。"
        return preview_msg

    print(structure_preview)

    # 9. 拆分标题段落：把"（一）xxx：正文"等拆成两段
    to_split_paragraphs = []
    for i, label in enumerate(labels):
        if label in ('heading1', 'heading2', 'heading3', 'heading4'):
            to_split_paragraphs.append((paragraphs[i], label))

    for para, label in to_split_paragraphs:
        result = split_heading_paragraph(para, label)
        if result is not None:
            changes['headings_split'].append((paragraphs.index(para), para.text[:30]))

    # 重新拉取段落列表
    paragraphs = body_paragraphs(doc)
    texts = [p.text for p in paragraphs]
    total = len(paragraphs)

    # 重新分类（因为段落数可能变化）
    labels = [classify_paragraph(t, i, total) for i, t in enumerate(texts)]
    # 重新修复：无编号标题检测（重分类后可能丢失）
    _fix_unnumbered_headings(paragraphs, labels)
    # 重新修复：连续附件检测（重分类后可能丢失）
    _fix_sequential_attachments(paragraphs, labels)
    # 重新识别标题（含多段标题合并）—— 首段空段时需要重新触发
    _ensure_title_paragraph(labels, texts)

    # 9.5 多段落标题合并：首段为标题时，将紧随其后的"纯标题续行"也并入标题，
    # 避免长标题被手工换行成多段后被误判为正文（如本行不以句号结尾且较短、且全文尚未出现一级标题）。
    if labels and labels[0] in ('title', 'title_candidate'):
        for j in range(1, len(labels)):
            lb = labels[j]
            if lb in ('heading1', 'heading2', 'heading3', 'heading4', 'recipient',
                      'attachment', 'date_signature', 'unit_signature', 'contact',
                      'doc_number', 'signer', 'banji', 'end_text', 'image_caption', 'empty'):
                break
            if lb == 'body':
                t = texts[j].strip()
                if (len(t) < 30 and not t.endswith(('。', '！', '？', '.', '．'))
                        and not any(l == 'heading1' for l in labels[:j])):
                    labels[j] = 'title'
                else:
                    break
            else:
                break

    # 9.6 层级编号修正（--fix-hierarchy）：在输出文件上修正跳跃的层级编号
    # 仅在 fix_hierarchy=True 时执行，只修改内存中的段落文本，不碰源文档
    if fix_hierarchy:
        hierarchy_fixes = fix_hierarchy_in_document(paragraphs, labels, texts)
        changes['hierarchy_fixes'] = hierarchy_fixes
        # 编号变更后重新分类（如 3、→（三）后类别从 heading3 变为 heading2）
        texts = [p.text for p in paragraphs]
        labels = [classify_paragraph(t, i, total) for i, t in enumerate(texts)]
        # 重新修复：无编号标题检测（重分类后可能丢失）
        _fix_unnumbered_headings(paragraphs, labels)
        # 重新修复：连续附件检测（重分类后可能丢失）
        _fix_sequential_attachments(paragraphs, labels)
        # 重新识别标题（hierarchy fix 可能改文本，需重做）
        _ensure_title_paragraph(labels, texts)

    # 9.6.5 后处理：智能识别落款（在所有段落重排/重分类后调用，避免索引错位）
    post_process_signatures(paragraphs, labels)

    # 9.7 附件计数与索引分配（用于多级悬挂缩进分发）
    # 按出现顺序为所有 attachment 段落分配 index，并统计总数
    attachment_total = sum(1 for lb in labels if lb == 'attachment')
    attachment_index_map = {}  # 段落索引 -> 附件序号（1-based）
    cur_index = 0
    for i, lb in enumerate(labels):
        if lb == 'attachment':
            cur_index += 1
            attachment_index_map[i] = cur_index

    # 10. 应用格式
    for i, paragraph in enumerate(paragraphs):
        label = labels[i]
        if label in ('title', 'title_candidate'):
            apply_title_format(paragraph, break_pos=title_break_pos)
            if title_break_pos is not None and title_break_pos > 0 and title_break_pos < len(paragraph.text or ''):
                # 找到对应的断行标签
                tb_info = diagnosis.get('title_break', (False, '', []))
                if tb_info[0]:
                    for pos, blabel, _ in tb_info[2]:
                        if pos == title_break_pos:
                            changes['title_broken'] = (pos, blabel)
                            break
                    if changes['title_broken'] is None:
                        changes['title_broken'] = (title_break_pos, '指定位置')
        elif label == 'recipient':
            apply_recipient_format(paragraph)
        elif label == 'heading1':
            apply_heading1_format(paragraph)
        elif label == 'heading2':
            apply_heading2_format(paragraph)
        elif label == 'heading3':
            apply_heading3_format(paragraph)
        elif label == 'heading4':
            apply_heading4_format(paragraph)
        elif label == 'attachment':
            idx = attachment_index_map.get(i, 1)
            apply_attachment_format(paragraph, attachment_index=idx, total_attachments=attachment_total)
        elif label == 'date_signature':
            apply_signature_format(paragraph)
        elif label == 'unit_signature':
            apply_signature_unit_format(paragraph)
        elif label == 'contact':
            apply_contact_format(paragraph)
        elif label == 'end_text':
            apply_end_text_format(paragraph)
        elif label == 'doc_number':
            apply_doc_number_format(paragraph)
        elif label == 'signer':
            apply_signer_format(paragraph)
        elif label == 'banji':
            apply_banji_format(paragraph)
        elif label == 'image_caption':
            apply_image_caption_format(paragraph)
        elif label == 'body':
            apply_body_format(paragraph)
        elif label == 'bullet_heading':
            # —— 开头段落未转换：回退为正文格式（保持现状）
            apply_body_format(paragraph)
            changes['bullet_headings_kept'].append((i, paragraph.text.strip()[:30]))
        # empty → 已在预处理中删除，此处不出现

    # 11. 表格排版（顶层段落已处理，表格单元格不在此循环内）
    for table in doc.tables:
        apply_table_format(table)
    changes['tables_formatted'] = len(doc.tables)

    # 12. 保存
    doc.save(output_path)

    # 13. 打印修改摘要
    summary = generate_change_summary(diagnosis, changes)
    print(summary)

    return output_path


def is_doc_file(filepath):
    """检测文件是否为旧版 .doc 格式（Compound Document File）"""
    try:
        with open(filepath, 'rb') as f:
            magic = f.read(8)
            # OLE2 / Compound Document File magic bytes
            return magic[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'
    except Exception:
        return False


def main():
    if len(sys.argv) < 2:
        print("用法: python format_gongwen.py <输入文件.docx> [输出文件.docx] [--preview] [--title-break=<pos>] [--fix-hierarchy] [--auto]")
        print("示例: python format_gongwen.py 报告草稿.docx 报告定稿.docx")
        print("      python format_gongwen.py 报告草稿.docx --preview  # 仅预览结构")
        print("      python format_gongwen.py 报告草稿.docx --title-break=12  # 标题第12字处断行")
        print("      python format_gongwen.py 报告草稿.docx --fix-hierarchy  # 排版+层级编号修正")
        print("      python format_gongwen.py 报告草稿.docx --auto  # 自动模式：跳过所有确认环节")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = None
    preview_only = False
    title_break_pos = None
    auto_mode = False
    fix_hierarchy = False

    for arg in sys.argv[2:]:
        if arg == '--preview':
            preview_only = True
        elif arg == '--auto':
            auto_mode = True
        elif arg == '--fix-hierarchy':
            fix_hierarchy = True
        elif arg.startswith('--title-break='):
            try:
                title_break_pos = int(arg.split('=', 1)[1])
            except ValueError:
                print(f"错误: --title-break 需要整数参数，收到: {arg}")
                sys.exit(1)
        elif not arg.startswith('--'):
            output_path = arg

    if not os.path.exists(input_path):
        print(f"错误: 文件不存在 - {input_path}")
        sys.exit(1)

    # 检测 .doc 旧格式
    if is_doc_file(input_path):
        print("=" * 50)
        print("⚠ 检测到旧版 .doc 格式（即使扩展名为 .docx）")
        print("  本工具基于 python-docx，仅支持 .docx (OOXML) 格式。")
        print("  请先用 Word 或 WPS 将文档另存为 .docx 格式后再排版。")
        print("=" * 50)
        sys.exit(1)

    print(f"正在排版: {input_path}")
    result = format_document(input_path, output_path, preview_only, title_break_pos, auto_mode, fix_hierarchy)
    if preview_only:
        print(result)
    else:
        print(f"排版完成: {result}")


if __name__ == "__main__":
    main()
